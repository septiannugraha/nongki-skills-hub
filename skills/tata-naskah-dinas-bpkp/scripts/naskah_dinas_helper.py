# -*- coding: utf-8 -*-
"""
Naskah Dinas Helper - BPKP
==========================

Modul pembantu untuk menyusun komponen Tata Naskah Dinas BPKP sesuai
Peraturan BPKP Nomor 4 Tahun 2022. Mencakup: kop surat, nota dinas,
surat tugas, dan lembar pengesahan.

Modul ini mengimpor engine inti (``bpkp_docx_engine``) untuk konsistensi
format (font, margin, heading styles). Ia menyediakan builder tingkat
tinggi untuk komponen naskah dinas yang sering dipakai berulang.
"""

from __future__ import annotations

from typing import Optional, List

import os
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Impor engine inti - relatif jika dipanggil sebagai paket, absolut jika mandiri.
# Engine bpkp_docx_engine berada di skill laporan-pengawasan-bpkp.
import os as _os
import sys as _sys

try:
    from .bpkp_docx_engine import (
        _apply_font, FONT_NAME, BLACK, add_p, add_run,
        clean_cell_p, set_cell_margins, set_cell_shading,
        set_table_borders, add_table_bordered,
    )
except ImportError:
    # Cari engine di skill laporan-pengawasan-bpkp (sibling directory)
    _here = _os.path.dirname(_os.path.abspath(__file__))
    _engine_dir = _os.path.normpath(
        _os.path.join(_here, '..', '..', 'laporan-pengawasan-bpkp', 'scripts')
    )
    if _engine_dir not in _sys.path:
        _sys.path.insert(0, _engine_dir)
    from bpkp_docx_engine import (
        _apply_font, FONT_NAME, BLACK, RED, add_p, add_run,
        clean_cell_p, set_cell_margins, set_cell_shading,
        set_cell_bottom_border, set_table_borders, add_table_bordered,
    )

__all__ = [
    "add_kop_surat",
    "add_kop_surat_table",
    "add_surat_pengantar_metadata",
    "add_nota_dinas_header",
    "add_surat_tugas_header",
    "add_lembar_pengesahan",
    "add_tte_marker",
    "add_table_bordered",
    "add_p",
    "add_run",
]

# =====================================================================
# KOP SURAT (LETTERHEAD) BPKP
# =====================================================================

def add_kop_surat(doc, unit_kerja: str, alamat: str = "", telepon: str = "",
                  email: str = "", website: str = "",
                  kode_pos: str = "", lembaga: str = "BADAN PENGAWASAN KEUANGAN DAN PEMBANGUNAN"):
    """
    Tambahkan kop surat (letterhead) BPKP versi paragraf sederhana.

    Untuk kop surat resmi dengan logo dan garis bawah tebal, gunakan
    fungsi ``add_kop_surat_table``.

    Layout:
      Baris 1: LEMBAGA (huruf besar, bold 14pt, tengah)
      Baris 2: UNIT KERJA (huruf besar, bold 12pt, tengah)
      Baris 3: Alamat lengkap (italic 10pt, tengah)
      Garis pembatas tebal (bottom border sz=18)
    """
    # Baris 1: Nama Lembaga
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(lembaga)
    _apply_font(r1, bold=True, size=Pt(14))

    # Baris 2: Unit Kerja
    if unit_kerja:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(unit_kerja.upper())
        _apply_font(r2, bold=True, size=Pt(12))

    # Baris 3: Alamat
    alamat_parts = [alamat]
    if telepon:
        alamat_parts.append(f"Telp. {telepon}")
    if email:
        alamat_parts.append(f"Email: {email}")
    if website:
        alamat_parts.append(f"Website: {website}")
    alamat_line = " | ".join(part for part in alamat_parts if part)
    if kode_pos:
        alamat_line = f"{alamat_line} - Kode Pos {kode_pos}" if alamat_line else f"Kode Pos {kode_pos}"

    if alamat_line:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(alamat_line)
        _apply_font(r3, italic=True, size=Pt(10))

    # Garis pembatas tebal (bottom border pada paragraf kosong)
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(6)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="18" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# =====================================================================
# KOP SURAT TABLE (dengan Logo + Garis Bawah)
# =====================================================================

def add_kop_surat_table(doc, logo_path: str = "",
                        lembaga: str = "BADAN PENGAWASAN KEUANGAN DAN PEMBANGUNAN",
                        unit_kerja: str = "PERWAKILAN PROVINSI PAPUA TENGAH",
                        alamat: str = "Jalan Sam Ratulangi, Kelurahan Oyehe, Distrik Nabire",
                        wilayah: str = "Kabupaten Nabire, Provinsi Papua Tengah, Kode Pos: 98816",
                        kontak: str = "Email: papua.tengah@bpkp.go.id, Website: www.bpkp.go.id/papua.tengah"):
    """
    Tambahkan kop surat (letterhead) BPKP versi tabel resmi dengan logo.

    Layout (tabel 1 baris x 2 kolom, borderless, dengan border bawah tebal):
      Kolom kiri  (3.2 cm): Logo BPKP (center)
      Kolom kanan (12.8 cm): 5 baris teks instansi
        - Baris 1: NAMA LEMBAGA       (bold 12pt, center, line 1.0)
        - Baris 2: UNIT KERJA          (bold 12pt, center, line 1.0)
        - Baris 3: Alamat jalan        (regular 10pt, center, line 1.0)
        - Baris 4: Wilayah + kode pos  (regular 10pt, center, line 1.0)
        - Baris 5: Email + Website     (regular 10pt, center, line 1.0)
      Garis bawah solid tebal (sz=18, hitam) pada kedua sel.
    """
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Borderless tabel, kecuali bottom border tebal di level tabel
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="single" w:sz="18" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Kolom kiri: Logo
    cell_logo = table.cell(0, 0)
    cell_logo.width = Cm(3.2)
    set_cell_margins(cell_logo, top=40, bottom=80, left=40, right=40)
    set_cell_bottom_border(cell_logo, color="000000", sz="18")

    p_klogo = cell_logo.paragraphs[0]
    p_klogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_klogo.paragraph_format.space_before = Pt(0)
    p_klogo.paragraph_format.space_after = Pt(2)
    clean_cell_p(p_klogo)
    if logo_path and os.path.exists(logo_path):
        r_klogo = p_klogo.add_run()
        r_klogo.add_picture(logo_path, width=Cm(2.9))

    # Kolom kanan: Teks instansi
    cell_ktxt = table.cell(0, 1)
    cell_ktxt.width = Cm(12.8)
    set_cell_margins(cell_ktxt, top=30, bottom=80, left=40, right=40)
    set_cell_bottom_border(cell_ktxt, color="000000", sz="18")

    lines = [
        {"text": lembaga, "bold": True, "size": 12, "sa": 1},
        {"text": unit_kerja, "bold": True, "size": 12, "sa": 3},
        {"text": alamat, "bold": False, "size": 10, "sa": 1},
        {"text": wilayah, "bold": False, "size": 10, "sa": 1},
        {"text": kontak, "bold": False, "size": 10, "sa": 4},
    ]
    for i, line in enumerate(lines):
        if i == 0:
            p = cell_ktxt.paragraphs[0]
        else:
            p = cell_ktxt.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(line["sa"])
        p.paragraph_format.line_spacing = 1.0
        clean_cell_p(p)
        add_run(p, line["text"], bold=line["bold"], size=Pt(line["size"]))

    # Ultra-thin separator paragraph
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after = Pt(4)
    pPr_sep = p_sep._p.get_or_add_pPr()
    pPr_sep.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="0" w:after="80" w:line="40" w:lineRule="exact"/>'
    ))
    r_sep = p_sep.add_run()
    rPr_sep = r_sep._r.get_or_add_rPr()
    rPr_sep.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="2"/>'))
    rPr_sep.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="2"/>'))


# =====================================================================
# TABEL METADATA SURAT PENGANTAR (4 Kolom)
# =====================================================================

def add_surat_pengantar_metadata(doc, nomor: str, lampiran: str, hal: str,
                                 tanggal: str = ""):
    """
    Tambahkan tabel metadata surat pengantar (4 kolom) standar BPKP.

    Layout (tabel 3 baris x 4 kolom, borderless):
      Kolom 1 (2.43 cm): Label (Nomor, Lampiran, Hal)      — left, 11pt
      Kolom 2 (0.63 cm): Separator ':'                      — center, 11pt
      Kolom 3 (8.39 cm): Isi teks                            — justified, 11pt
      Kolom 4 (4.58 cm): Tanggal (hanya baris pertama)       — right, 11pt

    Aturan: Tanggal ditempatkan pada kolom ke-4 mandiri agar judul
    laporan pada baris 'Hal' tidak mengalami hanging indent.
    """
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Borderless
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/><w:left w:val="none"/>'
        f'<w:right w:val="none"/><w:bottom w:val="none"/>'
        f'<w:insideH w:val="none"/><w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    meta_w = [Cm(2.43), Cm(0.63), Cm(8.39), Cm(4.58)]
    meta_rows = [
        ("Nomor", ":", nomor, tanggal),
        ("Lampiran", ":", lampiran, ""),
        ("Hal", ":", hal, ""),
    ]

    for r_idx, (label, sep, content, date_val) in enumerate(meta_rows):
        cells = table.rows[r_idx].cells
        for c_idx, w in enumerate(meta_w):
            cells[c_idx].width = w
            set_cell_margins(cells[c_idx], top=20, bottom=20, left=10, right=10)

        # Col 0: Label
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(2)
        p0.paragraph_format.line_spacing = 1.15
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        clean_cell_p(p0)
        add_run(p0, label, bold=False, size=Pt(11))

        # Col 1: Separator ':'
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.line_spacing = 1.15
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clean_cell_p(p1)
        add_run(p1, sep, bold=False, size=Pt(11))

        # Col 2: Content
        p2 = cells[2].paragraphs[0]
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.15
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        clean_cell_p(p2)
        add_run(p2, content, bold=False, size=Pt(11))

        # Col 3: Tanggal (kanan)
        p3 = cells[3].paragraphs[0]
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(2)
        p3.paragraph_format.line_spacing = 1.15
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        clean_cell_p(p3)
        if date_val:
            add_run(p3, date_val, bold=False, size=Pt(11))


# =====================================================================
# NOTA DINAS - HEADER
# =====================================================================

def add_nota_dinas_header(doc, nomor: str, sifat: str = "Biasa",
                          lampiran: str = "-", hal: str = "",
                          kepada: str = "", dari: str = "",
                          tempat_tanggal: str = ""):
    """
    Tambahkan header nota dinas standar BPKP.

    Layout (tabel 2 kolom):
      Kiri: Nomor, Sifat, Lampiran, Hal
      Kanan: Kepada (yth), Dari, Tempat/Tanggal
    """
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF", val="none")  # borderless

    # Lebar kolom: kiri 7 cm, kanan 9 cm
    for row in table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(9)

    labels_left = ["Nomor", "Sifat", "Lampiran", "Hal"]
    values_left = [nomor, sifat, lampiran, hal]
    labels_right = ["Kepada Yth.", "", "Dari", "Tempat/Tanggal"]
    values_right = [kepada, "", dari, tempat_tanggal]

    for i in range(4):
        # Kolom kiri
        cell_l = table.rows[i].cells[0]
        set_cell_margins(cell_l, top=40, bottom=40, left=0, right=100)
        p_l = cell_l.paragraphs[0]
        clean_cell_p(p_l)
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after = Pt(0)
        add_run(p_l, f"{labels_left[i]} : ", bold=False, size=Pt(12))
        add_run(p_l, values_left[i], bold=False, size=Pt(12))

        # Kolom kanan
        cell_r = table.rows[i].cells[1]
        set_cell_margins(cell_r, top=40, bottom=40, left=100, right=0)
        p_r = cell_r.paragraphs[0]
        clean_cell_p(p_r)
        p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_r.paragraph_format.space_before = Pt(2)
        p_r.paragraph_format.space_after = Pt(0)
        if labels_right[i]:
            add_run(p_r, f"{labels_right[i]} : ", bold=False, size=Pt(12))
        add_run(p_r, values_right[i], bold=False, size=Pt(12))

    # Baris terakhir: garis pemisah
    cell_sep = table.rows[4].cells[0]
    # Merge seluruh baris terakhir
    cell_sep.merge(table.rows[4].cells[1])
    set_cell_margins(cell_sep, top=0, bottom=0, left=0, right=0)
    p_sep = cell_sep.paragraphs[0]
    clean_cell_p(p_sep)
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after = Pt(0)
    pPr_sep = p_sep._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr_sep.append(pBdr)


# =====================================================================
# SURAT TUGAS - HEADER
# =====================================================================

def add_surat_tugas_header(doc, nomor: str, sifat: str = "Biasa",
                           lampiran: str = "-", hal: str = "Surat Tugas"):
    """
    Tambahkan header surat tugas standar BPKP.
    Mirip dengan nota dinas namun lebih sederhana (tanpa Dari/Kepada).
    """
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF", val="none")

    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12.5)

    labels = ["Nomor", "Sifat", "Lampiran", "Hal"]
    values = [nomor, sifat, lampiran, hal]

    for i in range(4):
        cell_l = table.rows[i].cells[0]
        set_cell_margins(cell_l, top=40, bottom=40, left=0, right=100)
        p_l = cell_l.paragraphs[0]
        clean_cell_p(p_l)
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after = Pt(0)
        add_run(p_l, f"{labels[i]} : ", bold=False, size=Pt(12))
        add_run(p_l, values[i], bold=False, size=Pt(12))

        # Kolom kanan kosong (untuk surat tugas, info pejabat ada di body)
        cell_r = table.rows[i].cells[1]
        set_cell_margins(cell_r, top=40, bottom=40, left=0, right=0)
        p_r = cell_r.paragraphs[0]
        clean_cell_p(p_r)
        add_run(p_r, "", size=Pt(12))

    # Garis pemisah
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after = Pt(6)
    pPr_sep = p_sep._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr_sep.append(pBdr)


# =====================================================================
# LEMBAR PENGESAHAN
# =====================================================================

def add_lembar_pengesahan(doc, judul_laporan: str, tempat_tanggal: str = "",
                          pejabat_list: Optional[List[dict]] = None):
    """
    Tambahkan halaman Lembar Pengesahan standar BPKP.

    pejabat_list: list of dict dengan keys:
      - nama    : nama pejabat
      - jabatan : jabatan pejabat
      - nip     : NIP pejabat (opsional)
    """
    if pejabat_list is None:
        pejabat_list = []

    # Judul
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("LEMBAR PENGESAHAN")
    _apply_font(r_title, bold=True, size=Pt(14))

    # Judul laporan
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run(judul_laporan)
    _apply_font(r_sub, bold=True, size=Pt(12))

    # Tabel pengesahan
    n_rows = max(len(pejabat_list), 1) + 1  # +1 untuk header
    table = doc.add_table(rows=n_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="000000", sz="6")

    col_widths = [Cm(6), Cm(7), Cm(4)]
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    # Header
    headers = ["Nama", "Jabatan", "Tanda Tangan"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_margins(hdr[i], top=100, bottom=100, left=100, right=100)
        set_cell_shading(hdr[i], "EAEAEA")
        p = hdr[i].paragraphs[0]
        clean_cell_p(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=Pt(11))

    # Isi
    for idx, pj in enumerate(pejabat_list, 1):
        row_cells = table.rows[idx].cells
        for c_idx in range(3):
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            clean_cell_p(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(row_cells[0].paragraphs[0], pj.get("nama", ""), size=Pt(11))
        add_run(row_cells[1].paragraphs[0], pj.get("jabatan", ""), size=Pt(11))
        # Kolom tanda tangan dibiarkan kosong
        add_run(row_cells[2].paragraphs[0], "", size=Pt(11))

    # Tempat/tanggal dan tanda tangan di bawah tabel
    if tempat_tanggal:
        add_p(doc, "", space_before=Pt(24))
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_date.paragraph_format.space_before = Pt(0)
        p_date.paragraph_format.space_after = Pt(0)
        r_date = p_date.add_run(tempat_tanggal)
        _apply_font(r_date, size=Pt(12))

        p_role = doc.add_paragraph()
        p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_role.paragraph_format.space_before = Pt(0)
        p_role.paragraph_format.space_after = Pt(48)  # space untuk TTE
        r_role = p_role.add_run(pejabat_list[0].get("jabatan", "") if pejabat_list else "")
        _apply_font(r_role, size=Pt(12))

        # TTE marker
        add_tte_marker(doc)

        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(0)
        p_name.paragraph_format.space_after = Pt(0)
        r_name = p_name.add_run(pejabat_list[0].get("nama", "") if pejabat_list else "")
        _apply_font(r_name, bold=True, size=Pt(12))

        p_nip = doc.add_paragraph()
        p_nip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nip.paragraph_format.space_before = Pt(0)
        p_nip.paragraph_format.space_after = Pt(0)
        nip_text = pejabat_list[0].get("nip", "")
        if nip_text:
            r_nip = p_nip.add_run(f"NIP. {nip_text}")
            _apply_font(r_nip, size=Pt(12))


# =====================================================================
# TTE MARKER (Tanda Tangan Elektronik)
# =====================================================================

def add_tte_marker(doc):
    """
    Tambahkan placeholder Tanda Tangan Elektronik (TTE).
    Diberi warna abu-abu italic sebagai indikasi visual.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("[Tanda Tangan Elektronik]")
    _apply_font(r, italic=True, color=RGBColor(89, 89, 89), size=Pt(10))


print("naskah_dinas_helper.py loaded - BPKP Tata Naskah Dinas helper ready.")
