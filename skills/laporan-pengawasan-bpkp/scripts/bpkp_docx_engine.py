# -*- coding: utf-8 -*-
"""
BPKP Document Engine
====================

Modul inti (engine) untuk pembangkitan dokumen Word (.docx) yang sesuai
dengan standar BPKP: A4, margin 3/2/2/2 cm, font Arial 12pt, spasi 1.15,
multilevel numbering A. -> 1. -> a. -> 1) -> a) dengan restart per Bab/Topik,
Word heading styles untuk panel navigasi, serta helper tabel dan blok
tanda tangan dinas.

Dirancang sebagai modul *reusable* - tidak bergantung pada data penugasan
spesifik manapun.  Setiap penugasan baru (LHP, LHE, LHR, evaluasi, reviu,
audit investigasi, dll.) cukup mengimpor fungsi-fungsi di sini.
"""

from __future__ import annotations

import os
from typing import Optional, List, Tuple

import docx
from docx import Document
from docx.document import Document as _Doc
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

__all__ = [
    "BLACK", "RED", "FONT_NAME",
    "create_document", "setup_heading_styles", "setup_numbering_infrastructure",
    "get_new_numbering_instance", "reset_numbering_state",
    "new_bab_context", "new_topic_context",
    "add_p", "add_run", "_apply_font",
    "add_heading_1", "add_heading_2", "add_heading_3",
    "add_heading_4",
    "add_section_heading", "add_topic_heading",
    "add_numbered_item", "add_simple_numbered",
    "add_subheading", "add_body_sub", "add_locus", "add_detail_item",
    "add_criteria", "add_cause", "add_effect",
    "add_recommendation", "add_recommendation_block",
    "clean_cell_p", "set_cell_margins", "set_cell_shading",
    "set_cell_bottom_border", "set_table_borders", "add_table_bordered",
    "add_signature_block",
    "add_cover_page",
    "add_daftar_isi",
    "add_page_break",
    "add_table_with_subheader",
    "_attach_numbering", "_LEVEL_FMT",
]

# =====================================================================
# KONSTANTA GLOBAL
# =====================================================================

BLACK = RGBColor(0, 0, 0)
RED = RGBColor(255, 0, 0)
FONT_NAME = 'Arial'

# =====================================================================
# STATE INTERNAL - melacak restart penomoran
# =====================================================================

_state: dict = {
    'current_bab': None,
    'current_topic': None,
    'current_bab_num_id': None,
}

# =====================================================================
# INFRASTRUKTUR NUMBERING (multilevel list OOXML)
#
# Satu abstractNum (ID 100) dengan 5 level terkait:
#   Level 0: A.   (Heading 2 - bagian seperti "A. Simpulan")
#   Level 1: 1.   (Heading 3 - judul topik / daftar bernomor)
#   Level 2: a.   (sub-temuan heading cetak tebal)
#   Level 3: 1)   (locus per wilayah)
#   Level 4: a)   (detail/rincian)
#
# Indentasi cascade: tiap level mulai di posisi teks level di atasnya.
#   L0: num@360  text@720
#   L1: num@720  text@1080
#   L2: num@1080 text@1440
#   L3: num@1440 text@1800
#   L4: num@1800 text@2160
# =====================================================================

_ABSTRACT_NUM_ID = 100
_next_num_id = 101

# (lvlText, left_twips, hanging_twips, tab_twips)
_LEVEL_FMT: List[Tuple[str, int, int, int]] = [
    ("%1.", 720,  360, 720),   # L0: A.
    ("%2.", 1080, 360, 1080),  # L1: 1.
    ("%3.", 1440, 360, 1440),  # L2: a.
    ("%4)", 1800, 360, 1800),  # L3: 1)
    ("%5)", 2160, 360, 2160),  # L4: a)
]


def _numfmt_name(ilvl: int) -> str:
    return {
        0: "upperLetter",
        1: "decimal",
        2: "lowerLetter",
        3: "decimal",
        4: "lowerLetter",
    }.get(ilvl, "decimal")


def _numfmt(ilvl: int) -> str:
    return {
        0: "%1.",
        1: "%2.",
        2: "%3.",
        3: "%4)",
        4: "%5)",
    }.get(ilvl, "%1.")


# =====================================================================
# BUILDER XML UNTUK NUMBERING
# =====================================================================

def _build_abstract_num_xml(abstract_id: int) -> "OxmlElement":
    """Bangun elemen <w:abstractNum> dengan 5 level terkait."""
    xml = (
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abstract_id}">'
        f'<w:multiLevelType w:val="multilevel"/>'
    )
    for i in range(5):
        nf = _numfmt_name(i)
        lt = _numfmt(i)
        _, left, hang, tab = _LEVEL_FMT[i]
        xml += (
            f'<w:lvl w:ilvl="{i}">'
            f'<w:start w:val="1"/>'
            f'<w:numFmt w:val="{nf}"/>'
            f'<w:lvlText w:val="{lt}"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:tabs><w:tab w:val="left" w:pos="{tab}"/></w:tabs>'
            f'<w:ind w:left="{left}" w:hanging="{hang}"/></w:pPr>'
            f'<w:rPr>'
            f'<w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
            f'w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
            f'<w:sz w:val="24"/><w:szCs w:val="24"/>'
            f'<w:color w:val="000000"/>'
            f'</w:rPr>'
            f'</w:lvl>'
        )
    xml += '</w:abstractNum>'
    return parse_xml(xml)


def _build_num_xml(num_id: int, abstract_id: int) -> "OxmlElement":
    """Bangun elemen <w:num> yang men-link numId ke abstractNumId."""
    return parse_xml(
        f'<w:num {nsdecls("w")} w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abstract_id}"/></w:num>'
    )


def _build_num_with_override(
    num_id: int, abstract_id: int, restart_levels: List[int], start: int = 1
) -> "OxmlElement":
    """Bangun <w:num> dengan lvlOverride untuk restart pada level tertentu."""
    xml = (
        f'<w:num {nsdecls("w")} w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abstract_id}"/>'
    )
    for ilvl in restart_levels:
        xml += (
            f'<w:lvlOverride w:ilvl="{ilvl}">'
            f'<w:startOverride w:val="{start}"/></w:lvlOverride>'
        )
    xml += '</w:num>'
    return parse_xml(xml)


# =====================================================================
# PENGELOLAAN NUMBERING PART
# =====================================================================

def _ensure_numbering_part(doc: _Doc):
    """Pastikan dokumen memiliki numbering part; buat jika belum ada."""
    try:
        doc.part.numbering_part
    except Exception:
        p = doc.add_paragraph(style='List Bullet')
        doc.part.numbering_part  # type: ignore
        p._element.getparent().remove(p._element)
    return doc.part.numbering_part


def _has_abstract_num(doc: _Doc, abstract_id: int) -> bool:
    try:
        for ab in doc.part.numbering_part.element.findall(qn('w:abstractNum')):
            if ab.get(qn('w:abstractNumId')) == str(abstract_id):
                return True
    except Exception:
        pass
    return False


def _has_num(doc: _Doc, num_id: int) -> bool:
    try:
        for n in doc.part.numbering_part.element.findall(qn('w:num')):
            if n.get(qn('w:numId')) == str(num_id):
                return True
    except Exception:
        pass
    return False


def setup_numbering_infrastructure(doc: _Doc) -> None:
    """Dipanggil sekali saat pembuatan dokumen untuk menyiapkan master abstractNum."""
    global _next_num_id
    _ensure_numbering_part(doc)
    np = doc.part.numbering_part
    elem = np.element

    if not _has_abstract_num(doc, _ABSTRACT_NUM_ID):
        ab = _build_abstract_num_xml(_ABSTRACT_NUM_ID)
        first_num = elem.find(qn('w:num'))
        if first_num is not None:
            elem.insert(list(elem).index(first_num), ab)
        else:
            elem.insert(0, ab)

    if not _has_num(doc, _ABSTRACT_NUM_ID):
        num = _build_num_xml(_ABSTRACT_NUM_ID, _ABSTRACT_NUM_ID)
        elem.append(num)


def get_new_numbering_instance(
    doc: _Doc, restart_levels: Optional[List[int]] = None, start_val: int = 1
) -> int:
    """
    Buat instance <w:num> baru yang menunjuk ke master abstractNum.
    Mengembalikan numId baru.
    """
    global _next_num_id
    np = doc.part.numbering_part
    nid = _next_num_id
    _next_num_id += 1

    if restart_levels:
        num_el = _build_num_with_override(nid, _ABSTRACT_NUM_ID, restart_levels, start_val)
    else:
        num_el = _build_num_xml(nid, _ABSTRACT_NUM_ID)

    np.element.append(num_el)
    return nid


def reset_numbering_state() -> None:
    """Reset state internal penomoran (mis. saat membuat dokumen baru)."""
    global _next_num_id
    _next_num_id = 101
    _state.clear()
    _state.update({
        'current_bab': None,
        'current_topic': None,
        'current_bab_num_id': None,
    })


# =====================================================================
# HELPER FONT & STYLE
# =====================================================================

def _apply_font(run, bold: bool = False, italic: bool = False,
               color: RGBColor = BLACK, size=Pt(12)) -> None:
    """Terapkan font Arial + atribut pada sebuah run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
        f'w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
    )
    rPr.append(rFonts)


def _ensure_style_font(style) -> None:
    """Pastikan sebuah paragraph style memakai Arial dan ukuran yang benar."""
    try:
        style.font.name = FONT_NAME
        style.font.size = Pt(12)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                f'w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
            )
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:ascii'), FONT_NAME)
            rFonts.set(qn('w:hAnsi'), FONT_NAME)
            rFonts.set(qn('w:cs'), FONT_NAME)
            rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except Exception:
        pass


def _set_style_spacing(style, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    """Atur spasi paragraf pada elemen style."""
    pPr = style.element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(
            f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}" '
            f'w:line="{int(line * 240)}" w:lineRule="auto"/>'
        )
        pPr.append(spacing)
    else:
        spacing.set(qn('w:before'), str(before))
        spacing.set(qn('w:after'), str(after))
        spacing.set(qn('w:line'), str(int(line * 240)))
        spacing.set(qn('w:lineRule'), 'auto')


def setup_heading_styles(doc: _Doc) -> None:
    """
    Konfigurasi style Heading 1, 2, 3 bawaan Word:
    - Font Arial 12pt bold hitam
    - Spasi before 12pt, after 6pt, line 1.15
    - Outline level untuk panel navigasi Word
    """
    for level, style_name in [(1, 'Heading 1'), (2, 'Heading 2'), (3, 'Heading 3'), (4, 'Heading 4')]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        _ensure_style_font(style)
        _set_style_spacing(style, before=240, after=120, line=1.15)
        style.font.bold = True
        style.font.size = Pt(12)
        style.font.color.rgb = BLACK
        pPr = style.element.get_or_add_pPr()
        outline = pPr.find(qn('w:outlineLvl'))
        if outline is None:
            outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level - 1}"/>')
            pPr.append(outline)
        else:
            outline.set(qn('w:val'), str(level - 1))


# =====================================================================
# CORE: create_document
# =====================================================================

def create_document() -> _Doc:
    """
    Buat dokumen Word kosong dengan setup standar BPKP:
    - Kertas A4 (21x29.7 cm)
    - Margin: kiri 3 cm, kanan 2 cm, atas 2 cm, bawah 2 cm
    - Font Arial 12pt, spasi 1.15 (default)
    - Heading styles untuk navigasi
    - Infrastruktur multilevel numbering
    """
    reset_numbering_state()

    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)

    normal = doc.styles['Normal']
    _ensure_style_font(normal)
    _set_style_spacing(normal, before=0, after=6, line=1.15)
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK

    setup_heading_styles(doc)
    setup_numbering_infrastructure(doc)
    return doc


# =====================================================================
# PARAGRAF & RUN BUILDERS
# =====================================================================

def add_p(doc, text="", space_before=Pt(0), space_after=Pt(6), line_spacing=1.15,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, left_indent=None, hanging_indent=None,
          tab_pos_dxa=None, bold=False, italic=False, color=BLACK, size=Pt(12),
          num_id=None, ilvl=None):
    """
    Tambahkan paragraf biasa dengan opsi penomoran multilevel.

    Jika num_id & ilvl diberikan, indentasi otomatis mengikuti posisi
    teks dari level induk.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    p.alignment = align
    if num_id is not None and ilvl is not None:
        fmt_data = _LEVEL_FMT[ilvl]
        left_inch = fmt_data[1] / 1440.0
        p.paragraph_format.left_indent = Inches(left_inch)
        p.paragraph_format.first_line_indent = Inches(0)
    else:
        if left_indent is not None:
            p.paragraph_format.left_indent = left_indent
        if hanging_indent is not None:
            p.paragraph_format.first_line_indent = hanging_indent
    if tab_pos_dxa is not None:
        pPr = p._p.get_or_add_pPr()
        tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="{tab_pos_dxa}"/></w:tabs>')
        pPr.append(tabs)
    if text:
        r = p.add_run(text)
        _apply_font(r, bold=bold, italic=italic, color=color, size=size)
    return p


def add_run(p, text, bold=False, italic=False, color=BLACK, size=Pt(12)):
    """Tambahkan run (segment teks) ke sebuah paragraf dengan font Arial."""
    r = p.add_run(text)
    _apply_font(r, bold=bold, italic=italic, color=color, size=size)
    return r


# =====================================================================
# HEADING FUNCTIONS - memakai Word built-in styles
# =====================================================================

def add_heading_1(doc, text):
    """Heading 1 - judul BAB. Memakai style 'Heading 1' -> muncul di Navigation Panel."""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _apply_font(r, bold=True, size=Pt(12))
    return p


def add_heading_2(doc, text, num_id=None):
    """
    Heading 2 - header bagian seperti 'A. Simpulan'.
    Memakai style 'Heading 2'. Jika num_id diberikan, lampirkan
    multilevel numbering level 0 (A.).
    """
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if num_id is not None:
        _attach_numbering(p, num_id, 0)
    r = p.add_run(text)
    _apply_font(r, bold=True, size=Pt(12))
    return p


def add_heading_3(doc, text, num_id=None):
    """
    Heading 3 - judul topik seperti '1. Lahan Pertanian...'.
    Memakai style 'Heading 3'. Jika num_id diberikan, lampirkan
    multilevel numbering level 1 (1.).
    """
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if num_id is not None:
        _attach_numbering(p, num_id, 1)
    r = p.add_run(text)
    _apply_font(r, bold=True, size=Pt(12))
    return p


def add_heading_4(doc, text, num_id=None):
    """
    Heading 4 - sub-temuan atau rincian poin spesifik seperti 'a. Pemenuhan Kuota...'.
    Memakai style 'Heading 4' -> muncul di Navigation Panel / Document Outline Word & Docs.
    Format baku: Arial 12 pt, Bold, Hitam (#000000).
    """
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if num_id is not None:
        _attach_numbering(p, num_id, 2)
    r = p.add_run(text)
    _apply_font(r, bold=True, size=Pt(12))
    return p


# =====================================================================
# NUMBERING ATTACHMENT
# =====================================================================

def _attach_numbering(paragraph, num_id, ilvl):
    """Lampirkan w:numPr ke pPr paragraf."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn('w:numPr'))
    if existing is not None:
        pPr.remove(existing)
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}><w:ilvl w:val="{ilvl}"/>'
        f'<w:numId w:val="{num_id}"/></w:numPr>'
    )
    pPr.append(numPr)


# =====================================================================
# NUMBERED PARAGRAPH BUILDERS
# =====================================================================

def add_section_heading(doc, text, num_id, ilvl=0, space_before=Pt(8), space_after=Pt(4)):
    """Heading cetak tebal dengan numbering multilevel (mis. A. / a.)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _attach_numbering(p, num_id, ilvl)
    r = p.add_run(text)
    _apply_font(r, bold=True, size=Pt(12))
    return p


def add_topic_heading(doc, text, num_id, ilvl=1, space_before=Pt(12), space_after=Pt(6)):
    """Heading topik (H3) dengan numbering level 1 (1.)."""
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _attach_numbering(p, num_id, ilvl)
    r = p.add_run(text)
    _apply_font(r, bold=True, size=Pt(12))
    return p


def add_numbered_item(doc, text, num_id, ilvl=2, bold=False,
                      space_before=Pt(0), space_after=Pt(6)):
    """Paragraf bernomor pada ilvl tertentu."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _attach_numbering(p, num_id, ilvl)
    r = p.add_run(text)
    _apply_font(r, bold=bold, size=Pt(12))
    return p


def add_simple_numbered(doc, text, num_id, ilvl=1, space_before=Pt(0), space_after=Pt(6)):
    """Tambah item bernomor sederhana memakai multilevel list."""
    return add_numbered_item(doc, text, num_id, ilvl=ilvl,
                             space_before=space_before, space_after=space_after)


# =====================================================================
# LEGACY COMPATIBLE BUILDERS (dipertahankan untuk kompatibilitas API lama)
# =====================================================================

def add_subheading(doc, prefix, text, num_id=None, ilvl=2):
    """
    Sub-temuan heading: 'a. Judul' (cetak tebal).
    Jika num_id diberikan, gunakan multilevel numbering level 2.
    Jika tidak, fallback ke prefix manual (legacy).
    """
    if num_id is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _attach_numbering(p, num_id, ilvl)
        r = p.add_run(text)
        _apply_font(r, bold=True, size=Pt(12))
        return p
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pPr = p._p.get_or_add_pPr()
        tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="504"/></w:tabs>')
        pPr.append(tabs)
        add_run(p, f"{prefix}\t", bold=True, color=BLACK)
        add_run(p, text, bold=True, color=BLACK)
        return p


def add_body_sub(doc, text, num_id=None, ilvl=None):
    """Paragraf isi di bawah subheading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if num_id is not None and ilvl is not None:
        fmt_data = _LEVEL_FMT[ilvl]
        left_inch = fmt_data[1] / 1440.0
        p.paragraph_format.left_indent = Inches(left_inch)
        p.paragraph_format.first_line_indent = Inches(0)
    else:
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(0.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, text, bold=False, color=BLACK)
    return p


def add_locus(doc, number_prefix, locus_name, narrative_text, num_id=None, ilvl=3):
    """
    Locus paragraph: '1) Provinsi Papua Tengah' sebagai sub-heading bernomor,
    diikuti narasi pada baris baru di bawahnya.
    """
    if num_id is not None:
        # Paragraf 1: nama locus bernomor
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _attach_numbering(p, num_id, ilvl)
        r = p.add_run(locus_name)
        _apply_font(r, bold=False, size=Pt(12))

        # Paragraf 2: narasi sejajar dengan teks ilvl
        fmt_data = _LEVEL_FMT[ilvl]
        left_inch = fmt_data[1] / 1440.0
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.line_spacing = 1.15
        p2.paragraph_format.left_indent = Inches(left_inch)
        p2.paragraph_format.first_line_indent = Inches(0)
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r2 = p2.add_run(narrative_text)
        _apply_font(r2, bold=False, size=Pt(12))
        return p2
    else:
        # Legacy
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.65)
        p.paragraph_format.first_line_indent = Inches(-0.30)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pPr = p._p.get_or_add_pPr()
        tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="936"/></w:tabs>')
        pPr.append(tabs)
        add_run(p, f"{number_prefix}\t", bold=False, color=BLACK)
        add_run(p, f"{locus_name}: ", bold=True, color=BLACK)
        add_run(p, narrative_text, bold=False, color=BLACK)
        return p


def add_detail_item(doc, item_prefix, item_text, bold_title="", num_id=None, ilvl=4):
    """Detail item: 'a) Judul: teks' (ilvl=4)."""
    if num_id is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _attach_numbering(p, num_id, ilvl)
        if bold_title:
            r = p.add_run(f"{bold_title}: ")
            _apply_font(r, bold=True, size=Pt(12))
        r2 = p.add_run(item_text)
        _apply_font(r2, bold=False, size=Pt(12))
        return p
    else:
        # Legacy
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.95)
        p.paragraph_format.first_line_indent = Inches(-0.30)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pPr = p._p.get_or_add_pPr()
        tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="1368"/></w:tabs>')
        pPr.append(tabs)
        add_run(p, f"{item_prefix}\t", bold=False, color=BLACK)
        if bold_title:
            add_run(p, f"{bold_title}: ", bold=True, color=BLACK)
        add_run(p, item_text, bold=False, color=BLACK)
        return p


def add_criteria(doc, criteria_text="(.)", num_id=None, ilvl=None):
    """Paragraf kriteria (merah)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if num_id is not None and ilvl is not None:
        fmt_data = _LEVEL_FMT[ilvl]
        left_inch = fmt_data[1] / 1440.0
        p.paragraph_format.left_indent = Inches(left_inch)
    else:
        p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(0.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, criteria_text, bold=False, color=RED)
    return p


def add_cause(doc, cause_text, num_id=None, ilvl=None):
    """Paragraf sebab."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if num_id is not None and ilvl is not None:
        fmt_data = _LEVEL_FMT[ilvl]
        left_inch = fmt_data[1] / 1440.0
        p.paragraph_format.left_indent = Inches(left_inch)
    else:
        p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(0.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, "Kondisi tersebut disebabkan oleh ", bold=False, color=BLACK)
    add_run(p, cause_text, bold=False, color=BLACK)
    return p


def add_effect(doc, effect_text, num_id=None, ilvl=None):
    """Paragraf akibat."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if num_id is not None and ilvl is not None:
        fmt_data = _LEVEL_FMT[ilvl]
        left_inch = fmt_data[1] / 1440.0
        p.paragraph_format.left_indent = Inches(left_inch)
    else:
        p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(0.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, "Akibatnya, ", bold=False, color=BLACK)
    add_run(p, effect_text, bold=False, color=BLACK)
    return p


def add_recommendation_block(doc, rec_points, num_id=None, ilvl=4):
    """Blok rekomendasi: intro + item bernomor (ilvl=4)."""
    parent_ilvl = max(ilvl - 1, 0) if ilvl > 0 else 0
    parent_fmt = _LEVEL_FMT[parent_ilvl]
    parent_text_pos = parent_fmt[1] / 1440.0

    if isinstance(rec_points, str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if num_id is not None:
            _attach_numbering(p, num_id, ilvl)
        else:
            pPr = p._p.get_or_add_pPr()
            ind_el = parse_xml(f'<w:ind {nsdecls("w")} w:left="360"/>')
            pPr.append(ind_el)
        add_run(p, "Atas permasalahan tersebut, direkomendasikan agar: ", bold=False, color=BLACK)
        add_run(p, rec_points, bold=False, color=BLACK)
        return p
    elif isinstance(rec_points, list):
        p_intro = doc.add_paragraph()
        p_intro.paragraph_format.space_before = Pt(0)
        p_intro.paragraph_format.space_after = Pt(3)
        p_intro.paragraph_format.line_spacing = 1.15
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if num_id is not None:
            p_intro.paragraph_format.left_indent = Inches(parent_text_pos)
        else:
            pPr = p_intro._p.get_or_add_pPr()
            ind_el = parse_xml(f'<w:ind {nsdecls("w")} w:left="360"/>')
            pPr.append(ind_el)
        add_run(p_intro, "Atas permasalahan tersebut, direkomendasikan agar:", bold=False, color=BLACK)

        for item in rec_points:
            add_numbered_item(doc, item, num_id, ilvl=ilvl)


def add_recommendation(doc, rec_text, num_id=None, ilvl=None):
    """Alias untuk add_recommendation_block."""
    return add_recommendation_block(doc, rec_text, num_id=num_id, ilvl=ilvl)


# =====================================================================
# HELPER TABEL & SEL
# =====================================================================

def clean_cell_p(p):
    """
    Pastikan indent kiri/kanan/firstLine = 0 di dalam sel tabel
    agar marker ruler tepat di awal kolom.
    """
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    pPr = p._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag.endswith('ind'):
            pPr.remove(child)
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="0" w:right="0" w:firstLine="0"/>'))


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Atur margin internal sel tabel (dalam dxa/twips)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/></w:tcMar>'
    )
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    """Beri shading (warna latar) pada sel tabel."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)


def set_cell_bottom_border(cell, color="000000", sz="18"):
    """Tambah border bawah saja pada sel (mis. untuk baris total)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:top w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_table_borders(table, color="B0B0B0", sz="4", val="single"):
    """
    Atur border tabel: hanya atas-bawah dan insideH (garis horizontal).
    Border kiri/kanan/insideV dihilangkan untuk tampilan clean BPKP.
    """
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_table_bordered(doc, rows, cols, col_widths=None, header_rows=1,
                       shade_header="EAEAEA", border_color="B0B0B0"):
    """
    Helper tingkat tinggi: buat tabel dengan border clean, header shading,
    dan lebar kolom otomatis.

    Mengembalikan tuple (table, rows_data) di mana rows_data adalah list
    dari list cell yang sudah dibuat - caller tinggal mengisi teks.
    """
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=border_color)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    rows_data = []
    for r_idx, row in enumerate(table.rows):
        row_cells = []
        for c_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            if r_idx < header_rows and shade_header:
                set_cell_shading(cell, shade_header)
            p = cell.paragraphs[0]
            clean_cell_p(p)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            row_cells.append(cell)
        rows_data.append(row_cells)
    return table, rows_data


# =====================================================================
# BLOK TANDA TANGAN DINAS
# =====================================================================

def add_signature_block(doc, date_text=None, title="Kepala Perwakilan,",
                        name="............................", left_dxa=5245):
    """
    Tambahkan blok tanda tangan dinas standar BPKP.

    date_text : teks tempat/tanggal (mis. "Nabire, 31 Agustus 2026")
    title     : jabatan pejabat penanda tangan
    name      : nama pejabat penanda tangan
    left_dxa  : indent kiri (twips) untuk blok tanda tangan
    """
    if date_text:
        p_date = doc.add_paragraph()
        pPr_d = p_date._p.get_or_add_pPr()
        pPr_d.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="{left_dxa}"/>'))
        pPr_d.append(parse_xml(
            f'<w:spacing {nsdecls("w")} w:before="120" w:after="0" '
            f'w:line="276" w:lineRule="auto"/>'
        ))
        r_date = p_date.add_run(date_text)
        _apply_font(r_date, size=Pt(12))

    p_title = doc.add_paragraph()
    pPr_t = p_title._p.get_or_add_pPr()
    pPr_t.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="{left_dxa}"/>'))
    pPr_t.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="0" w:line="276" w:lineRule="auto"/>'
    ))
    r_title = p_title.add_run(title)
    _apply_font(r_title, size=Pt(12))

    p_sp1 = doc.add_paragraph()
    pPr_sp1 = p_sp1._p.get_or_add_pPr()
    pPr_sp1.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="{left_dxa}"/>'))
    pPr_sp1.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="240" w:after="0" '
        f'w:line="276" w:lineRule="auto"/>'
    ))

    p_sub = doc.add_paragraph()
    pPr_s = p_sub._p.get_or_add_pPr()
    pPr_s.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="{left_dxa}"/>'))
    pPr_s.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="0" w:line="276" w:lineRule="auto"/>'
    ))
    r_sub = p_sub.add_run("Ditandatangani secara elektronik oleh")
    _apply_font(r_sub, italic=True, color=RGBColor(89, 89, 89), size=Pt(10))

    p_sp2 = doc.add_paragraph()
    pPr_sp2 = p_sp2._p.get_or_add_pPr()
    pPr_sp2.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="{left_dxa}"/>'))
    pPr_sp2.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="240" w:after="0" '
        f'w:line="276" w:lineRule="auto"/>'
    ))

    p_name = doc.add_paragraph()
    pPr_n = p_name._p.get_or_add_pPr()
    pPr_n.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="{left_dxa}"/>'))
    pPr_n.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="0" w:line="276" w:lineRule="auto"/>'
    ))
    r_name = p_name.add_run(name)
    _apply_font(r_name, size=Pt(12))


# =====================================================================
# NUMBERING CONTEXT MANAGER
# =====================================================================

def new_bab_context(doc, bab_label):
    """
    Panggil saat memulai Bab (chapter) baru.
    Mengembalikan num_id untuk konteks numbering baru.
    Restart SEMUA level (A., 1., a., 1), a)) ke 1.
    """
    _state['current_bab'] = bab_label
    _state['current_topic'] = None
    nid = get_new_numbering_instance(doc, restart_levels=[0, 1, 2, 3, 4])
    _state['current_bab_num_id'] = nid
    return nid


def new_topic_context(doc, topic_label):
    """
    Panggil saat memasuki topik baru (H3) dalam sebuah Bab.
    Mengembalikan num_id Bab saat ini.
    """
    _state['current_topic'] = topic_label
    return _state.get('current_bab_num_id')


# =====================================================================
# HELPER PATH ASET (LOGO BPKP)
# =====================================================================

def get_default_logo(variant: str = "png") -> str:
    """
    Kembalikan path absolut ke file logo BPKP yang dibundel di folder
    ``assets/`` skill ini.

    variant:
      "png" -> logo_bpkp.png  (untuk cover page, transparan)
      "jpg" -> logo_bpkp_kop.jpg (untuk kop surat tabel)

    Mengembalikan string kosong jika file tidak ditemukan.
    """
    _here = os.path.dirname(os.path.abspath(__file__))
    if variant == "jpg":
        path = os.path.join(_here, "..", "assets", "logo_bpkp_kop.jpg")
    else:
        path = os.path.join(_here, "..", "assets", "logo_bpkp.png")
    path = os.path.normpath(path)
    return path if os.path.exists(path) else ""


# =====================================================================
# COVER PAGE BUILDER
# =====================================================================

def add_page_break(doc):
    """Tambahkan page break ke dokumen."""
    doc.add_page_break()


def add_cover_page(doc, logo_path: str = "",
                    lembaga: str = "BADAN PENGAWASAN KEUANGAN DAN PEMBANGUNAN",
                    unit_kerja: str = "PERWAKILAN PROVINSI PAPUA TENGAH",
                    judul: str = "LAPORAN HASIL EVALUASI",
                    subjudul: str = "ATAS TATA KELOLA KETAHANAN PANGAN",
                    tahun: str = "TAHUN 2026",
                    nomor: str = "", tanggal: str = ""):
    """
    Bangun halaman cover (sampul) laporan BPKP standar.

    Layout:
      - Logo BPKP (center, 4.2 cm)
      - Nama Lembaga (bold 14pt, center)
      - Unit Kerja (bold 14pt, center)
      - Spasi tengah (3x paragraf kosong)
      - Judul laporan (bold 13pt, center)
      - Subjudul (bold 13pt, center)
      - Tahun (bold 13pt, center)
      - Spasi bawah (4x paragraf kosong)
      - Tabel Nomor/Tanggal (borderless, 2 baris x 3 kolom):
          [Label | : | Nilai]  (bold 11pt)
      - Page break
    """
    import os

    # Gunakan logo default jika tidak diberikan
    if not logo_path:
        logo_path = get_default_logo(variant="png")

    # Logo BPKP Center
    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(24)
        p_logo.paragraph_format.space_after = Pt(12)
        r_logo = p_logo.add_run()
        r_logo.add_picture(logo_path, width=Cm(4.2))

    # Header Nama Instansi
    add_p(doc, lembaga, space_before=Pt(0), space_after=Pt(2),
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(14))
    add_p(doc, unit_kerja, space_before=Pt(0), space_after=Pt(48),
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(14))

    # Spasi Tengah
    for _ in range(3):
        add_p(doc, "", space_after=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)

    # Judul Utama Laporan
    add_p(doc, judul, space_before=Pt(0), space_after=Pt(4),
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(13))
    add_p(doc, subjudul, space_before=Pt(0), space_after=Pt(4),
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(13))
    add_p(doc, tahun, space_before=Pt(0), space_after=Pt(48),
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(13))

    # Spasi Bawah
    for _ in range(4):
        add_p(doc, "", space_after=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)

    # Tabel NOMOR dan TANGGAL (borderless)
    if nomor or tanggal:
        tbl = doc.add_table(rows=2, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        tblPr = tbl._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none"/><w:left w:val="none"/>'
            f'<w:right w:val="none"/><w:bottom w:val="none"/>'
            f'<w:insideH w:val="none"/><w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        col_w = [Cm(3.0), Cm(0.5), Cm(8.0)]
        data_rows = [
            ("NOMOR", ":", nomor),
            ("TANGGAL", ":", tanggal),
        ]
        for r_idx, row_vals in enumerate(data_rows):
            row_cells = tbl.rows[r_idx].cells
            for c_idx, val in enumerate(row_vals):
                row_cells[c_idx].width = col_w[c_idx]
                set_cell_margins(row_cells[c_idx], top=20, bottom=20, left=10, right=10)
                p = row_cells[c_idx].paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                clean_cell_p(p)
                add_run(p, val, bold=True, size=Pt(11))

    doc.add_page_break()


# =====================================================================
# DAFTAR ISI BUILDER (dengan Dot Leaders)
# =====================================================================

def add_daftar_isi(doc, items: list):
    """
    Bangun Daftar Isi dengan dot leaders.

    items: list of (title, page) tuples.
      - Jika title mengandung 'BAB' atau 'RINGKASAN', cetak bold.
      - Page break setelah selesai.

    Contoh items:
      [("RINGKASAN EKSEKUTIF", "i"), ("BAB I  SIMPULAN DAN REKOMENDASI", "1"),
       ("   A.  Simpulan", "1"), ...]
    """
    add_heading_1(doc, "DAFTAR ISI")
    add_p(doc, "Halaman", space_after=Pt(12),
          align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)

    for title, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        is_bold = ('BAB' in title or 'RINGKASAN' in title)
        add_run(p, title, bold=is_bold)
        dots_len = max(5, 75 - len(title) - len(page))
        add_run(p, " " + "." * dots_len + " ")
        add_run(p, page, bold=is_bold)

    doc.add_page_break()


# =====================================================================
# TABEL DENGAN SUB-HEADER ROW
# =====================================================================

def add_table_with_subheader(doc, headers: list, sub_headers: list,
                              data_rows: list, col_widths=None,
                              header_size=Pt(10), sub_header_size=Pt(9),
                              body_size=Pt(10), border_color="B0B0B0"):
    """
    Bangun tabel dengan baris header utama + baris sub-header (mis. a/b/c/d/e).

    Layout (borderless-BPKP style, header shading abu-abu):
      Row 0: Header utama (bold, 10pt, center, shading EAEAEA)
      Row 1: Sub-header (italic, 9pt, center, shading F5F5F5)
      Row 2+: Data (regular, 10pt, left/center/right)

    headers     : list[str] - nama kolom header utama
    sub_headers : list[str] - sub-label kolom (mis. ["a", "b", "c", "d", "e = (c-d)"])
    data_rows   : list[list[str]] - data per baris
    col_widths  : list[Cm] - lebar tiap kolom
    """
    n_cols = len(headers)
    n_rows = 2 + len(data_rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=border_color)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    # Row 0: Header utama
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        set_cell_shading(hdr_cells[i], "EAEAEA")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        clean_cell_p(p)
        add_run(p, h, bold=True, size=header_size)

    # Row 1: Sub-header
    sub_cells = table.rows[1].cells
    for i, s in enumerate(sub_headers):
        set_cell_margins(sub_cells[i], top=60, bottom=60, left=100, right=100)
        set_cell_shading(sub_cells[i], "F5F5F5")
        p = sub_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        clean_cell_p(p)
        add_run(p, s, italic=True, size=sub_header_size)

    # Row 2+: Data
    for r_idx, row_vals in enumerate(data_rows, 2):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row_vals):
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            clean_cell_p(p)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, val, size=body_size)

    return table


# =====================================================================
# KOP SURAT (LETTERHEAD) TABEL RESMI BPKP
# =====================================================================

def add_kop_surat_table(doc, logo_path: str = "",
                        lembaga: str = "BADAN PENGAWASAN KEUANGAN DAN PEMBANGUNAN",
                        unit_kerja: str = "PERWAKILAN PROVINSI PAPUA TENGAH",
                        alamat: str = "Jalan Sam Ratulangi, Kelurahan Oyehe, Distrik Nabire",
                        wilayah: str = "Kabupaten Nabire, Provinsi Papua Tengah, Kode Pos: 98816",
                        kontak: str = "Email: papua.tengah@bpkp.go.id, Website: www.bpkp.go.id/papua.tengah"):
    """
    Tambahkan kop surat (letterhead) BPKP versi tabel resmi dengan logo.

    Format Standar:
      Tabel 1 baris x 2 kolom, borderless kecuali garis bawah tebal 2.25 pt (sz=18):
        - Kolom kiri  (3.06 cm / 86.8 pt): Logo BPKP resmi (Center, lebar 2.9 cm)
        - Kolom kanan (13.34 cm / 378.2 pt): 5 baris teks instansi
            1. NAMA LEMBAGA      (Arial 12pt Bold, center, line 1.0, spaceBelow 1pt)
            2. UNIT KERJA         (Arial 12pt Bold, center, line 1.0, spaceBelow 3pt)
            3. Alamat jalan       (Arial 10pt Regular, center, line 1.0, spaceBelow 1pt)
            4. Wilayah + kode pos (Arial 10pt Regular, center, line 1.0, spaceBelow 1pt)
            5. Email + Website    (Arial 10pt Regular, center, line 1.0, spaceBelow 4pt)
    """
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Borderless tabel, kecuali bottom border tebal di level tabel
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="single" w:sz="18" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Kolom kiri: Logo
    cell_logo = table.cell(0, 0)
    cell_logo.width = Cm(3.06)
    set_cell_margins(cell_logo, top=30, bottom=80, left=40, right=40)
    set_cell_bottom_border(cell_logo, color="000000", sz="18")

    p_klogo = cell_logo.paragraphs[0]
    p_klogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_klogo.paragraph_format.space_before = Pt(0)
    p_klogo.paragraph_format.space_after = Pt(2)
    clean_cell_p(p_klogo)
    
    if not logo_path:
        logo_path = get_default_logo(variant="png") or get_default_logo(variant="jpg")
    if logo_path and os.path.exists(logo_path):
        r_klogo = p_klogo.add_run()
        r_klogo.add_picture(logo_path, width=Cm(2.9))

    # Kolom kanan: Teks instansi
    cell_ktxt = table.cell(0, 1)
    cell_ktxt.width = Cm(13.34)
    set_cell_margins(cell_ktxt, top=30, bottom=80, left=40, right=40)
    set_cell_bottom_border(cell_ktxt, color="000000", sz="18")

    lines = [
        {"text": lembaga, "bold": True, "size": 12, "sa": 1},
        {"text": unit_kerja, "bold": True, "size": 12, "sa": 3},
        {"text": alamat, "bold": False, "size": 10, "sa": 1},
        {"text": wilayah, "bold": False, "size": 10, "sa": 1},
        {"text": kontak, "bold": False, "size": 10, "sa": 4},
    ]
    for i, line in enumerate(lines):
        if i == 0:
            p = cell_ktxt.paragraphs[0]
        else:
            p = cell_ktxt.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(line["sa"])
        p.paragraph_format.line_spacing = 1.0
        clean_cell_p(p)
        add_run(p, line["text"], bold=line["bold"], size=Pt(line["size"]))
