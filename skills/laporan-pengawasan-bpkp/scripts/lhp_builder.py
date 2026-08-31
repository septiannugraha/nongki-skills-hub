# -*- coding: utf-8 -*-
"""
LHP Builder - BPKP
==================

Modul builder tingkat tinggi untuk menyusun struktur laporan hasil
pengawasan (LHP/LHE/LHR) BPKP. Memanfaatkan engine inti
(``bpkp_docx_engine``) dan menyediakan kerangka:

- BAB I  : Simpulan dan Rekomendasi
- BAB II : Pendahuluan/Umum (latar belakang, tujuan, ruang lingkup,
           dasar hukum, kelembagaan)
- BAB III: Hasil Evaluasi/Pemeriksaan (struktur temuan 5C:
           Kondisi, Kriteria, Sebab, Akibat, Rekomendasi)

Modul ini HANYA menyusun kerangka struktur. Isi/substansi setiap
temuan harus diisi oleh pengguna (AI atau manusia) berdasarkan data
penugasan aktual.

Contoh penggunaan minimal::

    from bpkp_docx_engine import create_document, save_document
    from lhp_builder import build_bab_i_template, build_bab_ii_template

    doc = create_document()
    build_bab_i_template(doc, simpulan_list=[...], rekomendasi_list=[...])
    build_bab_ii_template(doc, ...)
    doc.save("LHP.docx")
"""

from __future__ import annotations

from typing import Optional, List, Tuple

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

try:
    from .bpkp_docx_engine import (
        BLACK, RED, FONT_NAME, add_p, add_run, add_heading_1, add_heading_2,
        add_heading_3, add_section_heading, add_topic_heading,
        add_numbered_item, add_simple_numbered, add_subheading,
        add_body_sub, add_locus, add_detail_item, add_criteria, add_cause,
        add_effect, add_recommendation, add_recommendation_block,
        add_signature_block, new_bab_context, new_topic_context,
        add_table_bordered, clean_cell_p, set_cell_margins, set_cell_shading,
    )
except ImportError:
    from bpkp_docx_engine import (
        BLACK, RED, FONT_NAME, add_p, add_run, add_heading_1, add_heading_2,
        add_heading_3, add_section_heading, add_topic_heading,
        add_numbered_item, add_simple_numbered, add_subheading,
        add_body_sub, add_locus, add_detail_item, add_criteria, add_cause,
        add_effect, add_recommendation, add_recommendation_block,
        add_signature_block, new_bab_context, new_topic_context,
        add_table_bordered, clean_cell_p, set_cell_margins, set_cell_shading,
    )

__all__ = [
    "build_bab_i_template",
    "build_bab_ii_template",
    "build_bab_iii_temuan",
    "build_temuan_5c",
]


# =====================================================================
# BAB I - SIMPULAN DAN REKOMENDASI
# =====================================================================

def build_bab_i_template(doc, simpulan_list: List[Tuple[str, str]],
                          rekomendasi_list: List[Tuple[str, List[str]]],
                          intro_text: str = ""):
    """
    Bangun BAB I (Simpulan dan Rekomendasi).

    simpulan_list    : list of (judul, isi) - masing-masing menjadi item
                       bernomor (1., 2., 3., ...) di bawah 'A. Simpulan'.
    rekomendasi_list : list of (judul, [list_rekomendasi]) - masing-masing
                       menjadi item bernomor di bawah 'B. Rekomendasi'.
    intro_text       : teks paragraf pembuka sebelum daftar simpulan.
    """
    nid = new_bab_context(doc, "BAB I")

    add_heading_1(doc, "BAB I")
    add_heading_1(doc, "SIMPULAN DAN REKOMENDASI")

    # A. SIMPULAN
    add_heading_2(doc, "Simpulan", num_id=nid)
    if intro_text:
        add_p(doc, intro_text, space_after=Pt(6), num_id=nid, ilvl=0)
    else:
        add_p(doc, "Berdasarkan hasil pengawasan, dapat disimpulkan hal-hal pokok sebagai berikut:",
              space_after=Pt(6), num_id=nid, ilvl=0)

    for judul, isi in simpulan_list:
        # Judul simpulan sebagai heading topik (1.)
        add_topic_heading(doc, judul, num_id=nid, ilvl=1)
        # Isi simpulan
        add_p(doc, isi, space_after=Pt(6), num_id=nid, ilvl=1)

    # B. REKOMENDASI
    add_heading_2(doc, "Rekomendasi", num_id=nid)
    add_p(doc, "Atas simpulan tersebut, direkomendasikan hal-hal sebagai berikut:",
          space_after=Pt(6), num_id=nid, ilvl=0)

    for judul, rec_list in rekomendasi_list:
        add_topic_heading(doc, judul, num_id=nid, ilvl=1)
        for rec in rec_list:
            add_numbered_item(doc, rec, num_id=nid, ilvl=2)


# =====================================================================
# BAB II - UMUM / PENDAHULUAN
# =====================================================================

def build_bab_ii_template(doc, latar_belakang: str = "",
                           tujuan: List[str] = None,
                           ruang_lingkup: List[str] = None,
                           dasar_hukum: List[str] = None,
                           kelembagaan: str = ""):
    """
    Bangun BAB II (Umum/Pendahuluan).

    Struktur:
      A. Latar Belakang
      B. Tujuan
      C. Ruang Lingkup
      D. Dasar Hukum
      E. Kelembagaan
    """
    nid = new_bab_context(doc, "BAB II")

    add_heading_1(doc, "BAB II")
    add_heading_1(doc, "UMUM")

    # A. Latar Belakang
    add_heading_2(doc, "Latar Belakang", num_id=nid)
    if latar_belakang:
        add_p(doc, latar_belakang, space_after=Pt(6), num_id=nid, ilvl=0)

    # B. Tujuan
    add_heading_2(doc, "Tujuan", num_id=nid)
    add_p(doc, "Pengawasan ini bertujuan untuk:", space_after=Pt(6), num_id=nid, ilvl=0)
    if tujuan:
        for item in tujuan:
            add_simple_numbered(doc, item, num_id=nid, ilvl=1)

    # C. Ruang Lingkup
    add_heading_2(doc, "Ruang Lingkup", num_id=nid)
    add_p(doc, "Ruang lingkup pengawasan meliputi:", space_after=Pt(6), num_id=nid, ilvl=0)
    if ruang_lingkup:
        for item in ruang_lingkup:
            add_simple_numbered(doc, item, num_id=nid, ilvl=1)

    # D. Dasar Hukum
    add_heading_2(doc, "Dasar Hukum", num_id=nid)
    add_p(doc, "Pengawasan ini dilaksanakan berdasarkan:", space_after=Pt(6), num_id=nid, ilvl=0)
    if dasar_hukum:
        for item in dasar_hukum:
            add_simple_numbered(doc, item, num_id=nid, ilvl=1)

    # E. Kelembagaan
    add_heading_2(doc, "Kelembagaan", num_id=nid)
    if kelembagaan:
        add_p(doc, kelembagaan, space_after=Pt(6), num_id=nid, ilvl=0)


# =====================================================================
# BAB III - HASIL EVALUASI/PEMERIKSAAN (TEMUAN 5C)
# =====================================================================

def build_bab_iii_temuan(doc, temuan_list: List[dict]):
    """
    Bangun BAB III (Hasil Evaluasi/Pemeriksaan) berisi temuan-temuan.

    Setiap temuan adalah dict dengan keys:
      - judul       : judul topik temuan (mis. "Tata Kelola Lahan...")
      - kondisi     : teks paragraf Kondisi
      - kriteria    : teks paragraf Kriteria (bisa "(.)" jika belum diverifikasi)
      - sebab       : teks paragraf Sebab
      - akibat      : teks paragraf Akibat
      - rekomendasi : list[str] daftar rekomendasi, atau str tunggal
    """
    nid = new_bab_context(doc, "BAB III")

    add_heading_1(doc, "BAB III")
    add_heading_1(doc, "HASIL EVALUASI")

    for temuan in temuan_list:
        build_temuan_5c(doc, temuan, num_id=nid)


def build_temuan_5c(doc, temuan: dict, num_id: int):
    """
    Bangun satu temuan dengan struktur 5C (Kondisi, Kriteria, Sebab,
    Akibat, Rekomendasi) memakai multilevel numbering.

    Menggunakan num_id dari konteks Bab aktif.
    """
    judul = temuan.get("judul", "Temuan")
    kondisi = temuan.get("kondisi", "")
    kriteria = temuan.get("kriteria", "(.)")
    sebab = temuan.get("sebab", "")
    akibat = temuan.get("akibat", "")
    rekomendasi = temuan.get("rekomendasi", [])

    # Judul topik temuan (1., 2., 3., ...)
    add_topic_heading(doc, judul, num_id=num_id, ilvl=1)

    # a. Kondisi
    add_subheading(doc, "a.", "Kondisi", num_id=num_id, ilvl=2)
    add_body_sub(doc, kondisi, num_id=num_id, ilvl=2)

    # b. Kriteria
    add_subheading(doc, "b.", "Kriteria", num_id=num_id, ilvl=2)
    add_body_sub(doc, kriteria, num_id=num_id, ilvl=2)

    # c. Sebab
    add_subheading(doc, "c.", "Sebab", num_id=num_id, ilvl=2)
    add_body_sub(doc, sebab, num_id=num_id, ilvl=2)

    # d. Akibat
    add_subheading(doc, "d.", "Akibat", num_id=num_id, ilvl=2)
    add_body_sub(doc, akibat, num_id=num_id, ilvl=2)

    # e. Rekomendasi
    add_subheading(doc, "e.", "Rekomendasi", num_id=num_id, ilvl=2)
    if isinstance(rekomendasi, str):
        add_recommendation_block(doc, rekomendasi, num_id=num_id, ilvl=4)
    elif isinstance(rekomendasi, list):
        add_recommendation_block(doc, rekomendasi, num_id=num_id, ilvl=4)


print("lhp_builder.py loaded - BPKP LHP/LHE builder template ready.")
